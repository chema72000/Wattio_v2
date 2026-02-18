"""Export diary markdown to .docx using python-docx."""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches


async def export_today(project_dir: Path) -> str:
    """Export today's diary to a .docx file."""
    return export_diary(project_dir, date.today().isoformat())


def export_diary(project_dir: Path, date_str: str) -> str:
    """Export a specific diary date to .docx.

    Args:
        project_dir: The project root directory.
        date_str: Date in YYYY-MM-DD format.

    Returns:
        Status message.
    """
    diary_dir = project_dir / "wattio" / "diary"
    md_file = diary_dir / f"{date_str}.md"

    if not md_file.exists():
        return f"No diary found for {date_str}."

    content = md_file.read_text(encoding="utf-8")
    docx_path = diary_dir / f"{date_str}.docx"

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for line in content.splitlines():
        line = line.rstrip()

        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("> "):
            # Blockquote — add as indented paragraph
            p = doc.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Inches(0.5)
        elif line.strip() == "---":
            doc.add_page_break()
        elif line.strip():
            doc.add_paragraph(line)

    doc.save(str(docx_path))
    return f"Exported to {docx_path.relative_to(project_dir)}"
