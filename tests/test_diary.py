"""Tests for the diary system."""

import struct
import zlib
from pathlib import Path

import pytest
from docx import Document

from wattio.diary.writer import DiaryWriter
from wattio.diary.export import export_diary


def _make_tiny_png(path: Path) -> None:
    """Create a minimal valid 1x1 white PNG file."""
    # IHDR
    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    ihdr_crc = struct.pack(">I", zlib.crc32(b"IHDR" + ihdr_data) & 0xFFFFFFFF)
    ihdr = struct.pack(">I", 13) + b"IHDR" + ihdr_data + ihdr_crc
    # IDAT — single white pixel (filter byte 0 + RGB)
    raw_data = zlib.compress(b"\x00\xff\xff\xff")
    idat_crc = struct.pack(">I", zlib.crc32(b"IDAT" + raw_data) & 0xFFFFFFFF)
    idat = struct.pack(">I", len(raw_data)) + b"IDAT" + raw_data + idat_crc
    # IEND
    iend_crc = struct.pack(">I", zlib.crc32(b"IEND") & 0xFFFFFFFF)
    iend = struct.pack(">I", 0) + b"IEND" + iend_crc

    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(b"\x89PNG\r\n\x1a\n" + ihdr + idat + iend)


class TestDiaryWriter:
    def test_creates_diary_file(self, tmp_project: Path) -> None:
        writer = DiaryWriter(tmp_project)
        writer.log_user("Hello Wattio")
        writer.close_session()

        diary_files = list((tmp_project / "wattio" / "diary").glob("*.md"))
        assert len(diary_files) == 1

        content = diary_files[0].read_text()
        assert "Wattio Session Diary" in content
        assert "Hello Wattio" in content
        assert "Session ended" in content

    def test_logs_tool_calls(self, tmp_project: Path) -> None:
        writer = DiaryWriter(tmp_project)
        writer.log_user("Find magnetics")
        writer.log_tool_call("magnetic_suggest", {"schematic_path": "test.asc"})
        writer.log_tool_result("Found 3 components")
        writer.log_assistant("Here are the results")
        writer.close_session()

        diary_files = list((tmp_project / "wattio" / "diary").glob("*.md"))
        content = diary_files[0].read_text()
        assert "magnetic_suggest" in content
        assert "Found 3 components" in content
        assert "Here are the results" in content

    def test_truncates_long_results(self, tmp_project: Path) -> None:
        writer = DiaryWriter(tmp_project)
        writer.log_user("test")
        writer.log_tool_result("x" * 5000)
        writer.close_session()

        diary_files = list((tmp_project / "wattio" / "diary").glob("*.md"))
        content = diary_files[0].read_text()
        assert "truncated" in content


class TestDiarySimulation:
    def test_log_simulation_run(self, tmp_project: Path) -> None:
        writer = DiaryWriter(tmp_project)
        writer.log_user("Run sim")
        writer.log_simulation(
            tool_name="ltspice_run",
            schematic="01 - LTspice/flyback/test.asc",
            summary="**Simulation complete:** `01 - LTspice/flyback/test.asc`",
            params={"load_resistance": "8"},
            traces=["V(OUT)", "I(L1)"],
        )
        writer.close_session()

        diary_files = list((tmp_project / "wattio" / "diary").glob("*.md"))
        content = diary_files[0].read_text()
        assert "#### " in content
        assert "\u2014 Simulation" in content
        assert "**Schematic:** `01 - LTspice/flyback/test.asc`" in content
        assert "`load_resistance=8`" in content
        assert "`V(OUT)`" in content
        assert "`I(L1)`" in content

    def test_log_simulation_with_plot(self, tmp_project: Path) -> None:
        writer = DiaryWriter(tmp_project)
        writer.log_user("Sweep")
        writer.log_simulation(
            tool_name="ltspice_sweep",
            schematic="01 - LTspice/flyback/test.asc",
            summary="**Sweep complete:** `load_resistance` from 4 to 12",
            params={"load_resistance": "4 \u2192 12 step 2"},
            traces=["V(OUT)"],
            plot_path="wattio/results/sweep_load_20260220_1405.png",
        )
        writer.close_session()

        diary_files = list((tmp_project / "wattio" / "diary").glob("*.md"))
        content = diary_files[0].read_text()
        assert "\u2014 Parameter sweep" in content
        assert "![Parameter sweep plot](wattio/results/sweep_load_20260220_1405.png)" in content

    def test_log_simulation_edit(self, tmp_project: Path) -> None:
        writer = DiaryWriter(tmp_project)
        writer.log_user("Edit schematic")
        writer.log_simulation(
            tool_name="ltspice_edit",
            schematic="01 - LTspice/flyback/test.asc",
            summary="Set **R1** = `10k` in working copy.",
            params={"component": "R1", "value": "10k"},
        )
        writer.close_session()

        diary_files = list((tmp_project / "wattio" / "diary").glob("*.md"))
        content = diary_files[0].read_text()
        assert "\u2014 Schematic edit" in content
        assert "`component=R1`" in content


class TestDiaryExport:
    def test_export_creates_docx(self, tmp_project: Path) -> None:
        # Create a diary markdown file
        diary_dir = tmp_project / "wattio" / "diary"
        md_file = diary_dir / "2026-01-01.md"
        md_file.write_text("# Wattio Session Diary — 2026-01-01\n\n## Session 14:00\n\nHello\n", encoding="utf-8")

        result = export_diary(tmp_project, "2026-01-01")
        assert "Exported" in result

        docx_file = diary_dir / "2026-01-01.docx"
        assert docx_file.exists()
        assert docx_file.stat().st_size > 0

    def test_export_missing_diary(self, tmp_project: Path) -> None:
        result = export_diary(tmp_project, "1999-01-01")
        assert "No diary found" in result

    def test_export_embeds_image(self, tmp_project: Path) -> None:
        # Create a tiny PNG in wattio/results/
        png_path = tmp_project / "wattio" / "results" / "sweep.png"
        _make_tiny_png(png_path)

        diary_dir = tmp_project / "wattio" / "diary"
        md_file = diary_dir / "2026-01-02.md"
        md_file.write_text(
            "# Diary\n\n"
            "#### 14:05 \u2014 Parameter sweep\n"
            "**Schematic:** `test.asc`\n\n"
            "![Sweep plot](wattio/results/sweep.png)\n",
            encoding="utf-8",
        )

        result = export_diary(tmp_project, "2026-01-02")
        assert "Exported" in result

        doc = Document(str(diary_dir / "2026-01-02.docx"))
        # Check that the document has an inline shape (the embedded image)
        inline_shapes = doc.inline_shapes
        assert len(inline_shapes) >= 1

    def test_export_handles_missing_image(self, tmp_project: Path) -> None:
        diary_dir = tmp_project / "wattio" / "diary"
        md_file = diary_dir / "2026-01-03.md"
        md_file.write_text(
            "# Diary\n\n"
            "![Missing plot](wattio/results/nonexistent.png)\n",
            encoding="utf-8",
        )

        # Should not crash
        result = export_diary(tmp_project, "2026-01-03")
        assert "Exported" in result

        doc = Document(str(diary_dir / "2026-01-03.docx"))
        # Fallback text should be present
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Image not found" in full_text
